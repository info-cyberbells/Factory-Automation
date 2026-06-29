import os
import io
import csv
import json
import base64
from fastapi import APIRouter, UploadFile, File, HTTPException
from pydantic import BaseModel, ConfigDict
import openai
import pypdf
import docx
import openpyxl

router = APIRouter(prefix="/ocr", tags=["AI Invoice Scanning"])

class InvoiceResponse(BaseModel):
    model_config = ConfigDict(protected_namespaces=())
    vendor_name: str
    invoice_number: str
    total_amount: float
    tax_amount: float
    raw_text: str

def safe_float(val):
    if not val:
        return 0.0
    try:
        clean_val = str(val).replace('$', '').replace('₹', '').replace(',', '').strip()
        return float(clean_val)
    except (ValueError, TypeError):
        return 0.0

def extract_text_from_file(file_bytes: bytes, filename: str) -> str:
    filename_lower = filename.lower()
    
    # 1. Plain Text
    if filename_lower.endswith(('.txt', '.log')):
        for encoding in ['utf-8', 'latin-1', 'utf-16']:
            try:
                return file_bytes.decode(encoding).strip()
            except Exception:
                continue
        return file_bytes.decode('utf-8', errors='ignore').strip()
        
    # 2. PDF
    elif filename_lower.endswith('.pdf'):
        try:
            pdf_file = io.BytesIO(file_bytes)
            reader = pypdf.PdfReader(pdf_file)
            text = []
            for page_num, page in enumerate(reader.pages):
                page_text = page.extract_text()
                if page_text:
                    text.append(f"--- Page {page_num + 1} ---\n{page_text}")
            return "\n\n".join(text).strip()
        except Exception as e:
            return f"[Error extracting text from PDF: {str(e)}]"
            
    # 3. Word DOCX
    elif filename_lower.endswith('.docx'):
        try:
            docx_file = io.BytesIO(file_bytes)
            doc = docx.Document(docx_file)
            text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        text.append(" | ".join(row_text))
            return "\n".join(text).strip()
        except Exception as e:
            return f"[Error extracting text from Word document: {str(e)}]"
            
    # 4. Excel XLSX
    elif filename_lower.endswith('.xlsx'):
        try:
            xlsx_file = io.BytesIO(file_bytes)
            wb = openpyxl.load_workbook(xlsx_file, data_only=True)
            text = []
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                text.append(f"--- Sheet: {sheet_name} ---")
                for row in sheet.iter_rows(values_only=True):
                    if any(cell is not None for cell in row):
                        row_str = " | ".join(str(cell) if cell is not None else "" for cell in row)
                        text.append(row_str)
            return "\n".join(text).strip()
        except Exception as e:
            return f"[Error extracting text from Excel document: {str(e)}]"
            
    # 5. CSV
    elif filename_lower.endswith('.csv'):
        try:
            text_content = file_bytes.decode('utf-8', errors='ignore')
            csv_file = io.StringIO(text_content)
            reader = csv.reader(csv_file)
            text = []
            for row in reader:
                if any(cell.strip() for cell in row):
                    text.append(" | ".join(row))
            return "\n".join(text).strip()
        except Exception as e:
            return f"[Error extracting text from CSV: {str(e)}]"
            
    return ""

@router.post("/scan-invoice", response_model=InvoiceResponse)
async def scan_invoice(file: UploadFile = File(...)):
    """
    Scans a purchase invoice using OpenAI Vision API or Text Extraction
    coupled with GPT-4o-mini for structured JSON parsing.
    """
    api_key = os.getenv("OPENAI_API_KEY")
    
    # Read the file bytes
    try:
        file_bytes = await file.read()
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Failed to read uploaded file: {str(e)}")
        
    filename = file.filename
    filename_lower = filename.lower()
    
    # System instructions for GPT-4o-mini structured JSON output
    system_prompt = (
        "You are an expert invoice extraction AI. Analyze the provided invoice content and extract these exact fields:\n"
        "1. vendor_name (vendor's business or company name issuing the invoice)\n"
        "2. invoice_number (invoice identifier/ID number)\n"
        "3. total_amount (the final total amount payable, as a number/float)\n"
        "4. tax_amount (the tax/VAT/GST amount, as a number/float)\n"
        "5. raw_text (a clean, beautiful, well-formatted markdown transcript/reconstruction of all the lines and tables in the invoice)\n\n"
        "Return ONLY a JSON object with these exact keys: 'vendor_name', 'invoice_number', 'total_amount', 'tax_amount', 'raw_text'."
    )
    
    if api_key and api_key != "your_openai_api_key_here":
        try:
            client = openai.OpenAI(api_key=api_key)
            
            # Determine if it's an image
            is_image = filename_lower.endswith(('.png', '.jpg', '.jpeg', '.webp'))
            
            if is_image:
                base64_image = base64.b64encode(file_bytes).decode('utf-8')
                mime_type = "image/jpeg"
                if filename_lower.endswith('.png'):
                    mime_type = "image/png"
                elif filename_lower.endswith('.webp'):
                    mime_type = "image/webp"
                    
                messages = [
                    {"role": "system", "content": system_prompt},
                    {
                        "role": "user",
                        "content": [
                            {"type": "text", "text": f"Extract details from this uploaded invoice image (Filename: {filename})."},
                            {
                                "type": "image_url",
                                "image_url": {
                                    "url": f"data:{mime_type};base64,{base64_image}"
                                }
                            }
                        ]
                    }
                ]
                raw_extracted = f"[Scanned Image: {filename}]"
            else:
                raw_extracted = extract_text_from_file(file_bytes, filename)
                messages = [
                    {"role": "system", "content": system_prompt},
                    {
                        "role": "user",
                        "content": f"Extract details from this invoice text content (Filename: {filename}):\n\n{raw_extracted}"
                    }
                ]
                
            response = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=messages,
                response_format={ "type": "json_object" }
            )
            
            data = json.loads(response.choices[0].message.content)
            
            # Form final raw text to store (prefer model-cleaned formatted markdown, fallback to extracted text)
            raw_text = data.get("raw_text") or raw_extracted or "No readable text content extracted"
            
            return InvoiceResponse(
                vendor_name=data.get("vendor_name", "Unknown Vendor"),
                invoice_number=data.get("invoice_number", "INV-000"),
                total_amount=safe_float(data.get("total_amount")),
                tax_amount=safe_float(data.get("tax_amount")),
                raw_text=raw_text
            )
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"OpenAI API Error: {str(e)}")
            
    else:
        # Fallback Mock Logic (For testing without hitting API limits or when key is missing)
        raw_extracted = extract_text_from_file(file_bytes, filename) if not filename_lower.endswith(('.png', '.jpg', '.jpeg', '.webp')) else ""
        mock_raw = raw_extracted or f"Mocked extraction for image: {filename}"
        return InvoiceResponse(
            vendor_name="Acme Nylon Suppliers (Mocked)",
            invoice_number=f"INV-{filename.split('.')[0][-4:]}",
            total_amount=45000.00,
            tax_amount=8100.00,
            raw_text=f"Extracted via Local Fallback Engine. (Add OPENAI_API_KEY for real extraction)\n\n--- Extracted Text ---\n{mock_raw}"
        )

